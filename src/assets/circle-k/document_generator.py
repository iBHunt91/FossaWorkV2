import logging
import sys
import json
from docxtpl import DocxTemplate, InlineImage
from io import BytesIO
import io
import os
import uuid
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
from PIL import Image
from docx import Document

# Update the premium blend ratio
PREMIUM_BLEND_RATIO = 0

def generate_document(data, template_path):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")
    
    try:
        # Set up basic logging to stdout
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
        
        logging.debug(f"generate_document function called with template_path: {template_path}")
        logging.debug(f"Template file exists: {os.path.exists(template_path)}")
        logging.debug(f"Template file size: {os.path.getsize(template_path) if os.path.exists(template_path) else 'N/A'}")
        
        doc = DocxTemplate(template_path)
        logging.debug("Template loaded successfully")
        
        if 'store_number' not in data:
            raise ValueError("Missing required field: store_number")
        
        # Process dispensers data
        processed_dispensers = []
        for dispenser in data.get("dispensers", []):
            for product in dispenser.get("products", []):
                if product != "Plus":  # Skip "Plus" for the first table
                    processed_dispensers.append({
                        "pump_no": dispenser["number"],
                        "product": product,
                        "leak_check": "Y",
                        "dispenser_type": "Gilbarco",
                        "dispenser_serial": "",
                        "allocation_mode": "Y"
                    })
        
        # Process data for the second table (organized by dispenser)
        second_table_rows = []
        for dispenser in data.get("dispensers", []):
            pump_numbers = dispenser["number"].split('/')
            for pump in pump_numbers:
                for product in dispenser.get("products", []):
                    second_table_rows.append({
                        "pump_no": pump.strip(),
                        "product": product,
                        "blend_ratio": "100%" if product == "Regular" else ("70%" if product == "Plus" else "0%"),
                        "meter_changed": "N",
                        "slow": "N/A"
                    })
        
        # Sort the second table rows by pump number and product order
        product_order = {"Regular": 1, "Plus": 2, "Premium": 3}
        second_table_rows.sort(key=lambda x: (int(x["pump_no"]), product_order.get(x["product"], 4)))

        # Remove any potential empty rows
        first_table_data = [row for row in processed_dispensers if any(row.values())]
        second_table_data = [row for row in second_table_rows if any(row.values())]

        context = {
            "VISIT_NUMBER": data.get("visit_number", ""),
            "STORE_NUMBER": data["store_number"],
            "ADDRESS": data.get("address", ""),
            "CITY": data.get("city", ""),
            "STATE": data.get("state", ""),
            "ZIP": data.get("zip", ""),
            "METER_TYPE": data.get("meter_type", ""),
            "TECHNICIAN_NAME": data.get("technician_name", ""),
            "dispensers": first_table_data,
            "second_table_rows": second_table_data,
            "has_dispensers": bool(first_table_data),
            "has_second_table": bool(second_table_data),
            "num_rows_first_table": len(first_table_data),
            "num_rows_second_table": len(second_table_data)
        }
        
        # Log each context value individually
        for key, value in context.items():
            logging.debug(f"Context {key}: {value}")

        logging.debug(f"Signature path in data: {data.get('signature_path')}")

        if 'signature_path' in data and os.path.exists(data['signature_path']):
            try:
                logging.debug(f"Processing signature from file: {data['signature_path']}")
                # Use an even smaller width for the signature
                signature_width = Mm(20)  # Reduced from Mm(30)
                context['signature'] = InlineImage(doc, data['signature_path'], width=signature_width)
                logging.debug(f"Signature image prepared successfully with width: {signature_width}")
            except Exception as img_error:
                logging.exception(f"Error preparing signature image: {str(img_error)}")
                context['signature'] = None
        else:
            logging.warning(f"No signature path provided or file not found: {data.get('signature_path')}")
            context['signature'] = None

        # Render the template with the context
        doc.render(context)

        # Remove the last row of each table if it's empty
        for table in doc.tables:
            if table.rows:
                last_row = table.rows[-1]
                if all(cell.text.strip() == '' for cell in last_row.cells):
                    table._element.remove(last_row._element)

        # Force signature insertion
        if context['signature']:
            for paragraph in doc.paragraphs:
                if '{{SIGNATURE}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{SIGNATURE}}', '')
                    run = paragraph.add_run()
                    # Use the same smaller width as defined earlier
                    run.add_picture(data['signature_path'], width=Mm(25))
                    logging.debug("Signature forcefully inserted with controlled size")
                    break
            else:
                logging.warning("{{SIGNATURE}} placeholder not found in the document")

        # Save the document to a bytes buffer
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Clean up the temporary signature file
        if 'signature_path' in data and os.path.exists(data['signature_path']):
            os.remove(data['signature_path'])
            logging.debug(f"Temporary signature file removed: {data['signature_path']}")

        return output

    except Exception as e:
        logging.exception("Error in generate_document function")
        raise

# Main function to run as a standalone script
def main():
    if len(sys.argv) < 3:
        print("Usage: python document_generator.py <template_path> <json_data>")
        sys.exit(1)
    
    try:
        template_path = sys.argv[1]
        json_data = sys.argv[2]
        
        # Parse the JSON data
        data = json.loads(json_data)
        
        # Generate the document
        output = generate_document(data, template_path)
        
        # Write binary content to stdout
        sys.stdout.buffer.write(output.getvalue())
        
    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()